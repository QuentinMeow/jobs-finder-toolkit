"""Focused tests for paragraph-based DOCX resume extraction."""

from __future__ import annotations

import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SCRIPT_DIR = Path(__file__).resolve().parents[1]
EXTRACT_SCRIPT = SCRIPT_DIR / "extract.py"
sys.path.insert(0, str(SCRIPT_DIR))

from extract import extract_to_yaml, extract_with_diagnostics  # noqa: E402


def _bold_paragraph(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.add_run(text).bold = True
    return paragraph


def _native_list_paragraph(document: Document, text: str):
    paragraph = document.add_paragraph(text)
    properties = paragraph._p.get_or_add_pPr()
    num_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_properties.append(level)
    num_properties.append(num_id)
    properties.append(num_properties)
    return paragraph


class ExtractResumeTests(unittest.TestCase):
    def setUp(self):
        self._temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(self._temp_dir.cleanup)
        self.temp_dir = Path(self._temp_dir.name)

    def _path(self, name: str = "resume.docx") -> Path:
        return self.temp_dir / name

    @staticmethod
    def _add_identity(document: Document) -> None:
        document.add_paragraph("Jordan Rivers")
        document.add_paragraph("Portland, OR • jordan.rivers@example.com • linkedin.com/in/jordanrivers")

    def _save_valid_shell(self, path: Path) -> Document:
        document = Document()
        self._add_identity(document)
        document.add_heading("Experience", level=1)
        document.add_paragraph(
            "Northwind Systems – Software Engineer  2020 – Present | Portland, OR"
        )
        document.add_paragraph("• Built reliable fictional services.")
        document.save(path)
        return document

    def test_extracts_two_employers_in_source_order(self):
        path = self._path()
        document = Document()
        self._add_identity(document)
        document.add_heading("Professional Experience", level=1)
        document.add_paragraph(
            "Northwind Systems – Senior Software Engineer  2022 – Present | Portland, OR"
        )
        document.add_paragraph("• Led a fictional platform migration.")
        document.add_paragraph("Contoso Labs")
        document.add_paragraph("Software Engineer")
        document.add_paragraph("2019 — 2022 | Seattle, WA")
        document.add_paragraph("• Improved a fictional deployment pipeline.")
        document.save(path)

        data = extract_to_yaml(str(path))

        self.assertNotIn("employer", data)
        self.assertEqual(
            [employer["company"] for employer in data["employers"]],
            ["Northwind Systems", "Contoso Labs"],
        )
        self.assertEqual(data["employers"][0]["role"], "Senior Software Engineer")
        self.assertEqual(data["employers"][1]["dates"], "2019 — 2022")
        self.assertEqual(
            data["employers"][1]["bullets"],
            ["Improved a fictional deployment pipeline."],
        )

    def test_preserves_promotions_as_repeated_company_entries(self):
        path = self._path()
        document = Document()
        self._add_identity(document)
        document.add_heading("Work History", level=1)
        _bold_paragraph(document, "Fabrikam Systems")
        _bold_paragraph(document, "Senior Platform Engineer")
        document.add_paragraph("2023 – Present | Remote")
        document.add_paragraph("• Led the senior-level fictional work.")
        document.add_paragraph("Platform Engineer  2021 – 2023 | Remote")
        document.add_paragraph("• Delivered the earlier fictional work.")
        document.save(path)

        result = extract_with_diagnostics(path)

        self.assertTrue(result.ok, result.diagnostics)
        assert result.data is not None
        employers = result.data["employers"]
        self.assertEqual(len(employers), 2)
        self.assertEqual([item["company"] for item in employers], ["Fabrikam Systems"] * 2)
        self.assertEqual(
            [item["role"] for item in employers],
            ["Senior Platform Engineer", "Platform Engineer"],
        )

    def test_detects_native_word_list_paragraphs(self):
        path = self._path()
        document = Document()
        self._add_identity(document)
        document.add_heading("Employment", level=1)
        document.add_paragraph(
            "Adventure Works - Site Reliability Engineer  2020 - Present | Austin, TX"
        )
        _native_list_paragraph(document, "Automated a fictional recovery workflow.")
        document.save(path)

        data = extract_to_yaml(str(path))

        self.assertEqual(
            data["employers"][0]["bullets"],
            ["Automated a fictional recovery workflow."],
        )

    def test_keeps_direct_bullets_and_extracts_explicit_projects(self):
        path = self._path()
        document = Document()
        self._add_identity(document)
        document.add_heading("Relevant Experience", level=1)
        document.add_paragraph(
            "Tailspin Toys – Backend Engineer  2021 – Present | Remote"
        )
        document.add_paragraph("• Improved service reliability across a fictional platform.")
        _bold_paragraph(document, "Project Atlas")
        _native_list_paragraph(document, "Built the fictional project ingestion path.")
        _native_list_paragraph(document, "Reduced the fictional project's processing delay.")
        document.save(path)

        data = extract_to_yaml(str(path))
        employer = data["employers"][0]

        self.assertEqual(
            employer["bullets"],
            ["Improved service reliability across a fictional platform."],
        )
        self.assertEqual(
            employer["projects"],
            [{
                "title": "Project Atlas",
                "bullets": [
                    "Built the fictional project ingestion path.",
                    "Reduced the fictional project's processing delay.",
                ],
            }],
        )

    def test_projects_and_experience_sections_preserve_source_order(self):
        path = self._path()
        document = Document()
        self._add_identity(document)
        document.add_heading("Projects", level=1)
        document.add_paragraph(
            "Lakemont University – Student Developer  2022 – 2024 | Lakemont, ST")
        _bold_paragraph(document, "Campus simulator")
        document.add_paragraph("• Built a deterministic fictional route simulator.")
        document.add_heading("Experience", level=1)
        document.add_paragraph(
            "Northwind Systems – Software Engineer  2024 – Present | Remote")
        document.add_paragraph("• Built a reliable fictional backend service.")
        document.save(path)

        data = extract_to_yaml(str(path))
        self.assertEqual(
            [employer["company"] for employer in data["employers"]],
            ["Lakemont University", "Northwind Systems"],
        )

    def test_decorative_drawing_warns_but_image_only_document_fails(self):
        decorated_path = self._path("decorated.docx")
        decorated = Document()
        self._add_identity(decorated)
        decorated.paragraphs[0].add_run()._r.append(OxmlElement("w:drawing"))
        decorated.add_heading("Experience", level=1)
        decorated.add_paragraph(
            "Northwind Systems – Software Engineer  2020 – Present | Remote")
        decorated.add_paragraph("• Built a reliable fictional backend service.")
        decorated.save(decorated_path)

        decorated_result = extract_with_diagnostics(decorated_path)
        self.assertTrue(decorated_result.ok, decorated_result.diagnostics)
        self.assertIn(
            "IGNORED_DECORATIVE_DRAWING",
            [diagnostic.code for diagnostic in decorated_result.diagnostics],
        )

        image_path = self._path("image-only.docx")
        image_only = Document()
        image_only.add_paragraph().add_run()._r.append(OxmlElement("w:drawing"))
        image_only.save(image_path)
        image_result = extract_with_diagnostics(image_path)
        self.assertFalse(image_result.ok)
        self.assertIn(
            "IMAGE_ONLY_DOCUMENT",
            [diagnostic.code for diagnostic in image_result.diagnostics],
        )

    def test_rejects_table_and_multi_column_layouts(self):
        table_path = self._path("table.docx")
        table_document = Document()
        self._add_identity(table_document)
        table_document.add_heading("Experience", level=1)
        table_document.add_table(rows=1, cols=2)
        table_document.save(table_path)

        columns_path = self._path("columns.docx")
        columns_document = Document()
        self._add_identity(columns_document)
        columns_document.add_heading("Experience", level=1)
        columns_document.add_paragraph(
            "Northwind Systems – Engineer  2020 – Present | Portland, OR"
        )
        columns_document.add_paragraph("• Built a fictional service.")
        columns = columns_document.sections[0]._sectPr.xpath("./w:cols")[0]
        columns.set(qn("w:num"), "2")
        columns_document.save(columns_path)

        cases = [
            (table_path, "UNSUPPORTED_TABLE_LAYOUT"),
            (columns_path, "UNSUPPORTED_MULTI_COLUMN_LAYOUT"),
        ]
        for path, code in cases:
            with self.subTest(code=code):
                result = extract_with_diagnostics(path)
                self.assertFalse(result.ok)
                self.assertIsNone(result.data)
                self.assertIn(code, [diagnostic.code for diagnostic in result.diagnostics])

    def test_empty_and_corrupt_files_return_actionable_diagnostics(self):
        empty_path = self._path("empty.docx")
        Document().save(empty_path)
        corrupt_path = self._path("corrupt.docx")
        corrupt_path.write_bytes(b"not a docx package")

        empty = extract_with_diagnostics(empty_path)
        corrupt = extract_with_diagnostics(corrupt_path)

        self.assertEqual(empty.diagnostics[0].code, "EMPTY_DOCUMENT")
        self.assertEqual(corrupt.diagnostics[0].code, "CORRUPT_DOCX")
        process = subprocess.run(
            [sys.executable, str(EXTRACT_SCRIPT), str(corrupt_path)],
            capture_output=True,
            text=True,
            check=False,
        )
        self.assertEqual(process.returncode, 2)
        self.assertIn("CORRUPT_DOCX", process.stderr)
        self.assertNotIn("Traceback", process.stderr)

    def test_accepts_alternate_heading_pipe_header_and_date_separators(self):
        path = self._path()
        document = Document()
        self._add_identity(document)
        document.add_heading("Employment History", level=1)
        document.add_paragraph(
            "Adventure Works | Site Reliability Engineer | Feb 2018 to Jul 2020 | Austin, TX"
        )
        document.add_paragraph("• Operated fictional production services.")
        document.add_paragraph(
            "Tailspin Toys − Platform Engineer  2020 − Present | Remote"
        )
        document.add_paragraph("• Built a fictional internal platform.")
        document.save(path)

        data = extract_to_yaml(str(path))

        self.assertEqual(len(data["employers"]), 2)
        self.assertEqual(data["employers"][0]["dates"], "Feb 2018 to Jul 2020")
        self.assertEqual(data["employers"][0]["location"], "Austin, TX")
        self.assertEqual(data["employers"][1]["dates"], "2020 − Present")
        self.assertEqual(data["employers"][1]["location"], "Remote")

    def test_ambiguous_experience_fails_instead_of_guessing(self):
        path = self._path()
        document = Document()
        self._add_identity(document)
        document.add_heading("Experience", level=1)
        document.add_paragraph("Northwind Systems")
        document.add_paragraph("This prose has no role or dates.")
        document.save(path)

        result = extract_with_diagnostics(path)

        self.assertFalse(result.ok)
        self.assertIn(
            "AMBIGUOUS_EXPERIENCE_STRUCTURE",
            [diagnostic.code for diagnostic in result.diagnostics],
        )


class ContactLineTests(unittest.TestCase):
    """Contact headers that are NOT the repo's own `@example.com` fixture shape.

    Every tracked fixture uses `City, ST • name@example.com • linkedin.com/...`,
    which is exactly what the old detector tested for — so nothing caught the four
    ordinary real-resume headers below returning '' with exit 0 and no warning.
    """

    def setUp(self):
        self._temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(self._temp_dir.cleanup)
        self.temp_dir = Path(self._temp_dir.name)

    def _extract(self, *preamble: str):
        path = self.temp_dir / f"resume-{abs(hash(preamble))}.docx"
        document = Document()
        document.add_paragraph("Jordan Rivers")
        for line in preamble:
            document.add_paragraph(line)
        document.add_heading("Experience", level=1)
        document.add_paragraph(
            "Northwind Systems – Software Engineer  2020 – Present | Portland, OR")
        document.add_paragraph("• Built reliable fictional services.")
        document.save(path)
        return extract_with_diagnostics(path)

    def test_recovers_the_four_shapes_that_used_to_be_dropped(self):
        for line in (
            "Springfield, IL | (555) 019-2837",                # phone, no email
            "City, State - Phone Number - E-mail",             # unfilled template
            "Austin, TX | jordanrivers.dev | 555-0100",        # personal domain
            "Austin, TX | github.com/jordanrivers | 555-0100",  # github, not linkedin
        ):
            with self.subTest(contact=line):
                result = self._extract(line)
                self.assertTrue(result.ok)
                self.assertEqual(result.data["contact_line"], line)
                self.assertEqual(result.diagnostics, [])

    def test_the_original_example_shape_still_works(self):
        line = "Portland, OR • jordan.rivers@example.com • linkedin.com/in/jordanrivers"
        result = self._extract(line)
        self.assertEqual(result.data["contact_line"], line)
        self.assertEqual(result.diagnostics, [])

    def test_unclassifiable_preamble_line_is_kept_verbatim_and_named(self):
        # The whole point: never return '' silently. A line the parser cannot
        # classify is preserved AND reported, so a human can see what happened.
        prose = ("Experienced backend engineer with eight years building "
                 "distributed systems.")
        result = self._extract(prose)
        self.assertTrue(result.ok)  # a warning, not a fatal
        self.assertEqual(result.data["contact_line"], prose)
        self.assertEqual([d.code for d in result.diagnostics],
                         ["UNRECOGNIZED_CONTACT_LINE"])
        self.assertIn(prose, result.diagnostics[0].message)

    def test_a_real_contact_line_wins_over_a_preceding_summary_line(self):
        # Position alone is not enough — the shape decides which candidate wins.
        contact = "Portland, OR | jordan.rivers@example.com | 555-0100"
        result = self._extract("Backend engineer and occasional woodworker.", contact)
        self.assertEqual(result.data["contact_line"], contact)
        self.assertEqual(result.diagnostics, [])

    def test_no_preamble_after_the_name_is_reported_not_silent(self):
        result = self._extract()
        self.assertTrue(result.ok)
        self.assertEqual(result.data["contact_line"], "")
        self.assertEqual([d.code for d in result.diagnostics],
                         ["MISSING_CONTACT_LINE"])

    def test_the_cli_prints_the_warning_and_still_exits_zero(self):
        # A warning nobody sees is not a warning: the CLI used to print
        # diagnostics only when extraction FAILED.
        path = self.temp_dir / "warn.docx"
        document = Document()
        document.add_paragraph("Jordan Rivers")
        document.add_paragraph("Backend engineer and occasional woodworker.")
        document.add_heading("Experience", level=1)
        document.add_paragraph(
            "Northwind Systems – Software Engineer  2020 – Present | Portland, OR")
        document.add_paragraph("• Built reliable fictional services.")
        document.save(path)

        process = subprocess.run(
            [sys.executable, str(EXTRACT_SCRIPT), str(path)],
            capture_output=True, text=True, check=False)
        self.assertEqual(process.returncode, 0, process.stderr)
        self.assertIn("[UNRECOGNIZED_CONTACT_LINE]", process.stderr)
        self.assertNotIn("Traceback", process.stderr)
        self.assertIn("contact_line:", process.stdout)  # YAML still on stdout

    def test_a_prose_paragraph_far_below_the_name_is_never_captured(self):
        # Only the first few preamble paragraphs are eligible, so a tagline block
        # cannot supply the contact line from four lines down.
        result = self._extract("Line one.", "Line two.", "Line three.",
                               "jordan.rivers@example.com")
        self.assertNotEqual(result.data["contact_line"], "jordan.rivers@example.com")
        self.assertEqual([d.code for d in result.diagnostics],
                         ["UNRECOGNIZED_CONTACT_LINE"])


if __name__ == "__main__":
    unittest.main()
